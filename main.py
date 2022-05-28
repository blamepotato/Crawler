from selenium import webdriver
import chromedriver_autoinstaller
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
import pickle

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from io import BytesIO
import requests
from docx.shared import Inches
import time
import traceback



# global variables
users = [] # the users you want to highlight

admin = [] # instructor of the website 

max_retry = 5 # how many times you can retry

def login_with_cookie(date):
    url = "redacted" + date
    # setup
    options = Options()
    options.add_argument("--disable-notifications")
    options.headless = True
    chromedriver_autoinstaller.install()
    driver = webdriver.Chrome(service=Service(), options=options)
    driver.get(url)
    set_cookie(driver)
    doc = Document()

    try:
        page_number = int(WebDriverWait(driver, 60).until(
            EC.visibility_of_element_located(
                (By.CLASS_NAME, "pageNum"))).text.split("/")[1])
        for i in range(page_number):
            try:
                WebDriverWait(driver, 120).until(EC.visibility_of_element_located((By.CLASS_NAME, "message-group")))
            except TimeoutException:
                global max_retry
                while max_retry > 0:
                    print("卡住了，我还能再试 " + str(max_retry) + " 次")
                    driver.quit()
                    max_retry -= 1
                    login_with_cookie(date)
                    return
                if max_retry <= 0:
                    print("寄，呼叫gy")
                    driver.quit()
                    return

            messages = driver.find_elements(By.CLASS_NAME, "message-group")
            while len(messages) != 20 and i != page_number - 1:
                print("trying")
                time.sleep(1)
                messages = driver.find_elements(By.CLASS_NAME, "message-group")
            for message in messages:
                author = message.find_element(By.CLASS_NAME, "nickName")
                content_wrapper = message.find_element(By.XPATH, ".//p | .//img[@class='img-style']")
                curr_time = message.find_element(By.CLASS_NAME, "time-style")
          
                if content_wrapper.text:
                    render_file(doc, author.text, curr_time.text,
                                content_wrapper.text, date)
                else:
                    if author.text not in admin:
                        continue
                    img_link = content_wrapper.get_attribute("src")
                    if not img_link:
                        print("teacher emote here", curr_time.text, author.text,
                              content_wrapper.text)
                        continue
                    add_image(doc, author.text, curr_time.text, img_link, date)
            click = driver.find_element(By.CLASS_NAME, "el-icon-arrow-right")
            driver.execute_script('arguments[0].click();', click)
        doc.save(date + ".docx")
        print("做完了，跑路！")
        driver.quit()
        return
    except Exception:
        print(traceback.format_exc())
        driver.quit()
        print("呼叫gy")
        return


def login_without_cookie():
    url = "redacted"
    # setup
    options = Options()
    options.add_argument("--disable-notifications")
    chromedriver_autoinstaller.install()
    driver = webdriver.Chrome(service=Service())
    driver.get(url)
    phone_number = "13616119638"
    enter_phone = driver.find_element(By.ID, "one")
    enter_phone.send_keys(phone_number)

    get_verification_code = driver.find_element(By.ID, "getCode")
    get_verification_code.click()
    verification_code = input("请输入验证码: \n")

    enter_code = driver.find_element(By.ID, "two")
    enter_code.send_keys(verification_code)

    login = driver.find_element(By.CLASS_NAME, "login")
    login.click()

    try:
        WebDriverWait(driver, 30).until(
            EC.presence_of_element_located((By.CLASS_NAME, "select_list"))
        )
        get_cookie(driver)

    except Exception as e:
        print(e)
        driver.quit()


def set_cookie(driver):
    cookies = pickle.load(open("cookies.pkl", "rb"))  # load_cookie
    for cookie in cookies:
        driver.add_cookie(cookie)


def get_cookie(driver):
    pickle.dump(driver.get_cookies(), open("cookies.pkl", "wb"))  # store_cookie


def render_file(doc, name, curr_time, content, doc_name):
    # credit: https://stackoverflow.com/questions/61801936/set-background-color-shading-on-normal-text-with-python-docx
    if name in users:
        color = 'FF7276'
    else:
        color = 'd3d3d3'
    doc.add_paragraph(name + "            " + curr_time)
    p = doc.add_paragraph()
    txt = content
    run = p.add_run(txt)
    tag = run._r
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    run.font.size = Pt(11)
    tag.rPr.append(shd)
    doc.save(doc_name + ".docx")


def add_image(doc, name, curr_time, img_link, doc_name):
    # https://stackoverflow.com/questions/24341589/python-docx-add-picture-from-the-web
    doc.add_paragraph(name + "            " + curr_time)
    try:
        response = requests.get(img_link)
    except:
        time.sleep(5)
        response = requests.get(img_link)
    binary_img = BytesIO(response.content)
    doc.add_picture(binary_img, width=Inches(2))
    doc.save(doc_name + ".docx")


def main():
    dates = input("请输入日期, 例子：20220401。可输入多个日期，用空格隔开，例子：20220401 20220402 20220403\n").split()
    for date in dates:
        login_with_cookie(date)


if __name__ == '__main__':
    main()



