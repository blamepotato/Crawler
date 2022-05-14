from selenium import webdriver
import chromedriver_autoinstaller
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
import pickle

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from io import BytesIO
import requests
from docx.shared import Inches
import time
import sys
from main import users
from main import admin


def login_with_cookie(date):
    url = "redacted" + date
    # setup
    options = Options()
    options.add_argument("--disable-notifications")
    options.headless = True
    chromedriver_autoinstaller.install()
    driver = webdriver.Chrome(service=Service(), options=options)
    driver.get(url)
    set_cookie(driver)
    doc = Document()

    try:
        page_number = int(WebDriverWait(driver, 60).until(EC.visibility_of_element_located((By.CLASS_NAME, "pageNum"))).text.split("/")[1])
        for i in range(page_number):
            WebDriverWait(driver, 60).until(
                EC.presence_of_element_located((By.CLASS_NAME, "message-group"))
            )
            authors = driver.find_elements(By.CLASS_NAME, "nickName")
            content_wrappers = driver.find_elements(By.CLASS_NAME, "ant-dropdown-trigger")
            times = driver.find_elements(By.CLASS_NAME, "time-style")
            assert(len(authors) == len(content_wrappers) == len(times))
            for author, content_wrapper, curr_time in zip(authors, content_wrappers, times):
                if content_wrapper.text:
                    content = content_wrapper.find_element(By.TAG_NAME, "p").text
                    render_file(doc, author.text, curr_time.text, content, date)
                else:
                    if author.text not in admin:
                        continue
                    img = content_wrapper.find_element(By.TAG_NAME, "img")
                    img_link = img.get_attribute("src")
                    add_image(doc, author.text, curr_time.text, img_link, date)
            click = driver.find_element(By.CLASS_NAME, "el-icon-arrow-right")
            driver.execute_script('arguments[0].click();', click)
        print("做完了，跑路！")
        driver.quit()
        return
    except Exception as e:
        print("error:")
        print(e)
        driver.quit()
        return





def set_cookie(driver):
    cookies = pickle.load(open("cookies.pkl", "rb"))  # load_cookie
    for cookie in cookies:
        driver.add_cookie(cookie)


def render_file(doc, name, curr_time, content, doc_name):
    # credit: https://stackoverflow.com/questions/61801936/set-background-color-shading-on-normal-text-with-python-docx
    if name in users:
        color = 'FF7276'
    else:
        color = 'd3d3d3'
    doc.add_paragraph(name + "            " + curr_time)
    p = doc.add_paragraph()
    txt = content
    run = p.add_run(txt)
    tag = run._r
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    run.font.size = Pt(11)
    tag.rPr.append(shd)
    doc.save(doc_name + ".docx")


def add_image(doc, name, curr_time, img_link, doc_name):
    # https://stackoverflow.com/questions/24341589/python-docx-add-picture-from-the-web
    doc.add_paragraph(name + "            " + curr_time)
    try:
        response = requests.get(img_link)
    except:
        time.sleep(5)
        response = requests.get(img_link)
    binary_img = BytesIO(response.content)
    doc.add_picture(binary_img, width=Inches(2))
    doc.save(doc_name + ".docx")


def main():
    date = sys.argv[1]
    login_with_cookie(date)


if __name__ == '__main__':
    main()




